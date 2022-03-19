import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE


os.chdir('C:/Users/John Ergen/Desktop/Personal/Job App Docs/Cover Letters/CustomizeCoverLetterApp/')
#print(os.getcwd())

#companyName = input('Company Name: ')

def ccl(companyName):
    templateFile = Document('Ergen Data Analyst Template Cover Letter.docx')
    newTitle = companyName + ' Ergen Data Analyst Cover Letter.docx'
        
    firstParagraph = templateFile.paragraphs[0]
    secondParagraph = templateFile.paragraphs[1]
        
    firstParText = firstParagraph.text
    replacedCompName = firstParText.replace('[COMPANY NAME]', companyName)
        
    newParg = secondParagraph.insert_paragraph_before(replacedCompName)
    newParg.style = templateFile.styles.add_style('Cover Letter Style', WD_STYLE_TYPE.PARAGRAPH)
        
    font = newParg.style.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)
        
    paragraph_format = newParg.paragraph_format
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.5
    paragraph_format.first_line_indent = Inches(.5)
        
    firstParagraph.clear()
    #print(firstParagraph.styles())
        
    templateFile.save(newTitle + '.docx')
    